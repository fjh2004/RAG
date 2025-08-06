from typing import List, Dict, Any
from pypdf import PdfReader
from pdf2image import convert_from_path  
import tempfile
import os
from docx import Document
import markdown
from paddleocr import PaddleOCR
import re
import logging
import comtypes.client  
from PIL import Image
from kg_processor import KGProcessor
import io

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        # 初始化OCR工具，支持中英文识别和角度校正
        self.ocr = PaddleOCR(use_angle_cls=True, lang='ch', show_log=False)
        # 初始化知识图谱处理器
        self.kg_processor = KGProcessor()
        
    def _post_process_document(self, document: Dict) -> Dict:
        """文档后处理：构建知识图谱"""
        try:
            self.kg_processor.build_kg(document)
            logger.info(f"知识图谱构建完成：{document['metadata']['source']}")
        except Exception as e:
            logger.error(f"知识图谱构建失败: {str(e)}")
        return document

    def load_document(self, file_path: str) -> Dict[str, Any]:
        """
        加载多种格式文档并返回原始内容
        支持：PDF、DOCX、DOC、Markdown
        """
        try:
            if file_path.endswith('.pdf'):
                return self._process_pdf(file_path)
            elif file_path.endswith('.docx'):
                return self._process_word_docx(file_path)
            elif file_path.endswith('.doc'):
                return self._process_word_doc(file_path)
            elif file_path.endswith('.md'):
                return self._process_markdown(file_path)
            elif file_path.endswith('.txt'):
                return self._process_txt(file_path)
            else:
                raise ValueError(f"不支持的文件格式: {os.path.splitext(file_path)[1]}")
        except Exception as e:
            logger.error(f"处理文件 {file_path} 时出错: {str(e)}")
            raise

    def chunk_document(self, content: Dict[str, Any], chunk_size: int = 1000) -> List[Dict[str, Any]]:
        """
        智能分块策略：结合语义分割和固定长度，保留标题层级信息
        """
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        
        # 根据文档类型使用不同的分隔符，优先保留标题结构
        if content['metadata']['type'] == 'markdown':
            # Markdown优先按标题级别分割
            separators = [
                "\n# ", "\n## ", "\n### ", "\n#### ", "\n##### ", "\n###### ",  # 标题分隔符
                "\n\n", "\n", "。", "！", "？", ". ", "! ", "? ", " "  # 常规分隔符
            ]
        else:
            # 其他文档按常规标点分隔
            separators = ["\n\n", "\n", "。", "！", "？", ". ", "! ", "? ", " "]
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=200,  # 重叠部分保留上下文
            separators=separators,
            length_function=len  # 按字符数计算长度
        )
        
        chunks = []
        for chunk in text_splitter.create_documents([content['text']]):
            chunk_metadata = {
                "source": content['metadata']['source'],
                "document_type": content['metadata']['type'],
                "chunk_id": len(chunks)  # 块编号，方便追踪
            }
            
            # 提取并保留标题信息（针对Markdown）
            if content['metadata']['type'] == 'markdown':
                # 查找块中最高级别的标题
                headings = re.findall(r'^(#{1,6})\s+(.*?)$', chunk.page_content, re.MULTILINE)
                if headings:
                    # 按标题级别排序（#数量越少级别越高）
                    headings.sort(key=lambda x: len(x[0]))
                    chunk_metadata['heading'] = headings[0][1]
                    chunk_metadata['heading_level'] = len(headings[0][0])
            
            chunks.append({
                "text": chunk.page_content,
                "metadata": chunk_metadata
            })
        
        return chunks
    
    # document_processor.py 中添加
    def _process_txt(self, file_path: str) -> Dict[str, Any]:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return self._post_process_document({
            "text": text,
            "metadata": {"source": file_path, "type": "txt"}
        })

    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        处理PDF文档：
        1. 优先提取文本内容
        2. 若文本为空（图片型PDF），自动转为图片进行OCR识别
        """
        reader = PdfReader(file_path)
        text = ""
        page_texts = []
        
        # 尝试提取文本
        for page in reader.pages:
            page_text = page.extract_text() or ""
            page_texts.append(page_text)
            text += page_text + "\n\n"  # 每页文本间加空行分隔
        
        # 检查是否为图片型PDF（所有页都无文本）
        is_image_based = not any(page_text.strip() for page_text in page_texts)
        if is_image_based:
            logger.info(f"PDF {file_path} 为图片型，正在进行OCR处理...")
            text = self._ocr_pdf_pages(file_path)
        
        return self._post_process_document({
            "text": text,
            "metadata": {
                "source": file_path,
                "type": "pdf",
                "pages": len(reader.pages),
                "is_image_based": is_image_based
            }
        })
    
    def _ocr_pdf_pages(self, file_path: str) -> str:
        """将PDF每页转为图片并进行OCR识别"""
        try:
            # 将PDF转换为图片（需要安装poppler-utils）
            pages = convert_from_path(file_path, dpi=300)  # 高DPI提高识别精度
            full_text = ""
            
            # 使用临时目录存储图片
            with tempfile.TemporaryDirectory() as temp_dir:
                for i, page in enumerate(pages):
                    # 保存为临时图片
                    temp_image_path = os.path.join(temp_dir, f"page_{i+1}.png")
                    page.save(temp_image_path, "PNG")
                    
                    # 调用OCR识别
                    ocr_result = self.process_image(temp_image_path)
                    full_text += f"【第{i+1}页】\n{ocr_result['text']}\n\n"
            
            return full_text
        except Exception as e:
            logger.error(f"PDF OCR处理失败: {str(e)}")
            raise

    def _process_word_docx(self, file_path: str) -> Dict[str, Any]:
        """
        处理Word docx文档：
        1. 提取段落文本
        2. 处理表格并转换为结构化文本
        3. 提取图片并进行OCR识别
        """
        doc = Document(file_path)
        content_parts = []
        
        # 1. 处理段落文本
        for para in doc.paragraphs:
            if para.text.strip():  # 跳过空段落
                content_parts.append(para.text)
        
        # 2. 处理表格
        for table_idx, table in enumerate(doc.tables):
            table_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            table_text = self.process_table(table_data)
            content_parts.append(f"【表格 {table_idx+1}】\n{table_text}")
        
        # 3. 处理图片并进行OCR
        image_count = 0
        with tempfile.TemporaryDirectory() as temp_dir:
            # 遍历文档中的关系（包含图片）
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
                    # 提取图片二进制数据
                    image_part = rel.target_part
                    image_bytes = image_part._blob
                    
                    # 保存为临时图片
                    img_ext = image_part.content_type.split('/')[-1]  # 获取图片格式
                    temp_image_path = os.path.join(temp_dir, f"image_{image_count}.{img_ext}")
                    with open(temp_image_path, 'wb') as f:
                        f.write(image_bytes)
                    
                    # OCR识别图片文字
                    ocr_result = self.process_image(temp_image_path)
                    content_parts.append(f"【图片 {image_count}】\n{ocr_result['text']}")
        
        return {
            "text": '\n\n'.join(content_parts),
            "metadata": {
                "source": file_path,
                "type": "docx",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "images": image_count
            }
        }
    
    def _process_word_doc(self, file_path: str) -> Dict[str, Any]:
        """处理Word doc格式（旧版二进制格式）"""
        try:
            # 使用comtypes调用Word程序（仅支持Windows系统）
            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False  # 后台运行，不显示窗口
            
            # 打开文档
            doc = word.Documents.Open(os.path.abspath(file_path))
            
            # 提取文本内容
            text = doc.Content.Text
            
            # 获取文档信息
            table_count = doc.Tables.Count
            
            # 关闭文档和Word程序
            doc.Close()
            word.Quit()
            
            return {
                "text": text,
                "metadata": {
                    "source": file_path,
                    "type": "doc",
                    "tables": table_count,
                    "note": "doc格式暂不支持图片提取，建议转为docx格式以获得更好体验"
                }
            }
        except Exception as e:
            logger.error(f"处理doc文件失败: {str(e)}")
            raise

    def process_image(self, image_path: str) -> Dict[str, Any]:
        """处理图片OCR识别，支持倾斜校正"""
        try:
            # 执行OCR识别
            result = self.ocr.ocr(image_path, cls=True)  # cls=True启用角度检测
            
            if not result or not result[0]:
                return {
                    "text": "",
                    "metadata": {
                        "source": image_path,
                        "type": "image",
                        "ocr_success": False
                    }
                }
            
            # 提取识别文本（过滤低置信度结果）
            text_lines = []
            for line in result[0]:
                text, confidence = line[1]
                if confidence > 0.5:  # 过滤置信度低于0.5的结果
                    text_lines.append(text)
            
            return {
                "text": '\n'.join(text_lines),
                "metadata": {
                    "source": image_path,
                    "type": "image",
                    "ocr_success": True,
                    "line_count": len(text_lines)
                }
            }
        except Exception as e:
            logger.error(f"图片OCR处理失败: {str(e)}")
            return {
                "text": "",
                "metadata": {
                    "source": image_path,
                    "type": "image",
                    "ocr_success": False
                }
            }

    def process_table(self, table_data: List[List[str]]) -> str:
        """将表格数据转换为结构化文本，保留行列关系"""
        if not table_data:
            return "空表格"
        
        # 计算每列最大宽度，用于对齐（优化可读性）
        col_widths = [max(len(cell) for cell in col) for col in zip(*table_data)]
        
        # 生成表格文本
        table_lines = []
        for row in table_data:
            # 格式化每行，按列宽对齐
            formatted_row = [cell.ljust(width) for cell, width in zip(row, col_widths)]
            table_lines.append("| ".join(formatted_row))
        
        return '\n'.join(table_lines)
    
    def _process_markdown(self, file_path: str) -> Dict[str, Any]:
        """处理Markdown文档，保留标题层级和结构"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # 保留原始Markdown格式（而非转换为HTML），确保分块时能识别标题
            text = md_content
            
            # 统计标题层级分布
            heading_counts = {
                'h1': len(re.findall(r'^#\s', md_content, re.MULTILINE)),
                'h2': len(re.findall(r'^##\s', md_content, re.MULTILINE)),
                'h3': len(re.findall(r'^###\s', md_content, re.MULTILINE))
            }
            
            return {
                "text": text,
                "metadata": {
                    "source": file_path,
                    "type": "markdown",
                    **heading_counts
                }
            }
        except Exception as e:
            logger.error(f"处理Markdown文件失败: {str(e)}")
            raise
